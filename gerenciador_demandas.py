import tkinter as tk
from tkinter import messagebox, scrolledtext
from tkinter import font as tkfont
from tkinter import ttk
from datetime import datetime
import os
import subprocess
from docx import Document

# Define o caminho do diretório onde os arquivos Word serão salvos
diretorio_demandas = r'C:\Programas\Metas\Demandas'
os.makedirs(diretorio_demandas, exist_ok=True)

def salvar_anotacao(status, nota, data, solicitante, ritm, nome_arquivo):
    try:
        doc = Document()
        doc.add_heading('Demanda', level=1)
        doc.add_paragraph(f"Status: {status}")
        doc.add_paragraph(f"Nota: {nota}")
        doc.add_paragraph(f"Data: {data}")
        doc.add_paragraph(f"Solicitante: {solicitante}")
        doc.add_paragraph(f"Número da RITM: {ritm}")
        doc.add_paragraph("-" * 40)
        
        caminho_arquivo = os.path.join(diretorio_demandas, nome_arquivo + '.docx')
        doc.save(caminho_arquivo)
        print(f"Arquivo salvo com sucesso: {caminho_arquivo}")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao salvar a anotação: {e}")

def enviar_dados():
    status = status_var.get()
    nota = nota_text.get("1.0", tk.END).strip()
    data = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    solicitante = solicitante_entry.get()
    ritm = ritm_entry.get()
    nome_arquivo = nome_arquivo_entry.get()

    if not status or not nota or not solicitante or not ritm or not nome_arquivo:
        messagebox.showwarning("Aviso", "Todos os campos devem ser preenchidos!")
        return

    salvar_anotacao(status, nota, data, solicitante, ritm, nome_arquivo)
    messagebox.showinfo("Informação", "Anotação salva com sucesso!")
    status_var.set('Pendente')
    nota_text.delete("1.0", tk.END)
    solicitante_entry.delete(0, tk.END)
    ritm_entry.delete(0, tk.END)
    nome_arquivo_entry.delete(0, tk.END)
    atualizar_projetos_salvos()

def buscar_demandas():
    texto_busca = busca_entry.get().strip()
    if not texto_busca:
        messagebox.showwarning("Aviso", "Digite o número da RITM para busca!")
        return

    resultados = []
    try:
        arquivos = [f for f in os.listdir(diretorio_demandas) if f.endswith('.docx')]
        for arquivo in arquivos:
            doc = Document(os.path.join(diretorio_demandas, arquivo))
            conteudo = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            if f"Número da RITM: {texto_busca}" in conteudo:
                resultados.append(conteudo)
    except FileNotFoundError:
        messagebox.showwarning("Aviso", "Nenhuma demanda encontrada!")
        return
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao buscar demandas: {e}")
        return

    resultado_text.delete("1.0", tk.END)
    if resultados:
        resultado_text.insert(tk.END, "\n\n".join(resultados))
    else:
        resultado_text.insert(tk.END, "Nenhuma demanda encontrada.")

def abrir_projeto(nome_projeto):
    caminho_arquivo = os.path.join(diretorio_demandas, nome_projeto + '.docx')
    if not os.path.exists(caminho_arquivo):
        messagebox.showwarning("Aviso", "Arquivo não encontrado!")
        return

    try:
        if os.name == 'nt':  # Windows
            os.startfile(caminho_arquivo)
        elif os.name == 'posix':  # Unix-like
            subprocess.call(['open', caminho_arquivo])  # Para macOS
        else:
            messagebox.showwarning("Aviso", "Sistema operacional não suportado para abrir arquivos.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao abrir o arquivo: {e}")

def marcar_como_entregue(nome_projeto):
    caminho_arquivo = os.path.join(diretorio_demandas, nome_projeto + '.docx')
    if not os.path.exists(caminho_arquivo):
        messagebox.showwarning("Aviso", "Arquivo não encontrado!")
        return

    try:
        doc = Document(caminho_arquivo)
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith("Status:"):
                paragraph.text = "Status: Entregue"
        
        doc.save(caminho_arquivo)
        messagebox.showinfo("Informação", "Projeto marcado como entregue!")
        atualizar_projetos_salvos()
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao atualizar o status: {e}")

def excluir_demanda(nome_projeto):
    caminho_arquivo = os.path.join(diretorio_demandas, nome_projeto + '.docx')
    if not os.path.exists(caminho_arquivo):
        messagebox.showwarning("Aviso", "Arquivo não encontrado!")
        return

    try:
        os.remove(caminho_arquivo)
        messagebox.showinfo("Informação", "Demanda excluída com sucesso!")
        atualizar_projetos_salvos()
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao excluir a demanda: {e}")

def obter_status(nome_projeto):
    caminho_arquivo = os.path.join(diretorio_demandas, nome_projeto + '.docx')
    if os.path.exists(caminho_arquivo):
        try:
            doc = Document(caminho_arquivo)
            for paragraph in doc.paragraphs:
                if paragraph.text.startswith("Status:"):
                    return paragraph.text.split(": ", 1)[1]
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao ler o status: {e}")
    return "Pendente"

def atualizar_projetos_salvos():
    for widget in projetos_frame.winfo_children():
        widget.destroy()

    arquivos = [f for f in os.listdir(diretorio_demandas) if f.endswith('.docx')]
    if arquivos:
        for arquivo in arquivos:
            nome_projeto = arquivo.replace('.docx', '')
            frame_projeto = tk.Frame(projetos_frame, bg=cor_aba, pady=10)
            frame_projeto.pack(fill=tk.X, padx=10, pady=5)

            # Botão de abrir projeto
            btn_abrir = tk.Button(frame_projeto, text="Abrir", command=lambda f=nome_projeto: abrir_projeto(f), font=font_padrao, bg=cor_botao, fg="white", relief=tk.RAISED, bd=2)
            btn_abrir.pack(side=tk.LEFT, padx=5)

            # Botão de marcar como entregue
            status_text = obter_status(nome_projeto)
            if status_text == "Entregue":
                btn_entregar = tk.Button(frame_projeto, text="Entregue", font=font_padrao, bg="red", fg="white", relief=tk.RAISED, bd=2, state=tk.DISABLED)
            else:
                btn_entregar = tk.Button(frame_projeto, text="Marcar como Entregue", command=lambda f=nome_projeto: marcar_como_entregue(f), font=font_padrao, bg="#FFC107", fg="black", relief=tk.RAISED, bd=2)
            btn_entregar.pack(side=tk.LEFT, padx=5)

            # Botão de excluir projeto
            btn_excluir = tk.Button(frame_projeto, text="Excluir", command=lambda f=nome_projeto: excluir_demanda(f), font=font_padrao, bg="#F44336", fg="white", relief=tk.RAISED, bd=2)
            btn_excluir.pack(side=tk.LEFT, padx=5)

            # Descrição do projeto
            descricao_frame = tk.Frame(frame_projeto, bg=cor_aba)
            descricao_frame.pack(padx=5, pady=5, fill=tk.X)

            tk.Label(descricao_frame, text=f"Descrição do Projeto {nome_projeto}:", font=font_padrao, bg=cor_aba, fg=cor_label).pack(padx=5, pady=5)

            descricao_texto = scrolledtext.ScrolledText(descricao_frame, height=4, width=80, wrap=tk.WORD, font=font_padrao, bg=cor_aba)
            descricao_texto.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            try:
                doc = Document(os.path.join(diretorio_demandas, nome_projeto + '.docx'))
                descricao_texto.insert(tk.END, "\n".join(paragraph.text for paragraph in doc.paragraphs))
            except Exception as e:
                descricao_texto.insert(tk.END, f"Erro ao carregar descrição: {e}")

# Configuração da interface gráfica
root = tk.Tk()
root.title("Gerenciador de Demandas - Developer by Alexandre Martins")
root.geometry("1000x700")

# Define as cores e fontes
font_padrao = tkfont.Font(family="Helvetica", size=12)
cor_aba = "#F0F0F0"
cor_label = "#333333"
cor_botao = "#2196F3"

# Cria as abas
abas = ttk.Notebook(root)
aba_novo_projeto = ttk.Frame(abas)
aba_projetos_salvos = ttk.Frame(abas)

abas.add(aba_novo_projeto, text="Novo Projeto")
abas.add(aba_projetos_salvos, text="Projetos Salvos")
abas.pack(expand=1, fill="both")

# Adiciona o conteúdo da aba "Novo Projeto"
form_frame = tk.Frame(aba_novo_projeto, padx=20, pady=20, bg=cor_aba)
form_frame.pack(fill=tk.BOTH, expand=True)

tk.Label(form_frame, text="Status:", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
status_var = tk.StringVar(value="Pendente")
tk.Radiobutton(form_frame, text="Pendente", variable=status_var, value="Pendente", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=0, column=1, padx=5, pady=5)
tk.Radiobutton(form_frame, text="Em Andamento", variable=status_var, value="Em Andamento", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=0, column=2, padx=5, pady=5)
tk.Radiobutton(form_frame, text="Entregue", variable=status_var, value="Entregue", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=0, column=3, padx=5, pady=5)

tk.Label(form_frame, text="Nota:", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
nota_text = scrolledtext.ScrolledText(form_frame, height=8, width=60, wrap=tk.WORD, font=font_padrao)
nota_text.grid(row=1, column=1, columnspan=3, padx=10, pady=5)

tk.Label(form_frame, text="Nome do Solicitante:", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
solicitante_entry = tk.Entry(form_frame, width=60, font=font_padrao)
solicitante_entry.grid(row=2, column=1, columnspan=3, padx=10, pady=5)

tk.Label(form_frame, text="Número da RITM:", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=3, column=0, sticky=tk.W, padx=10, pady=5)
ritm_entry = tk.Entry(form_frame, width=60, font=font_padrao)
ritm_entry.grid(row=3, column=1, columnspan=3, padx=10, pady=5)

tk.Label(form_frame, text="Nome do Arquivo:", font=font_padrao, bg=cor_aba, fg=cor_label).grid(row=4, column=0, sticky=tk.W, padx=10, pady=5)
nome_arquivo_entry = tk.Entry(form_frame, width=60, font=font_padrao)
nome_arquivo_entry.grid(row=4, column=1, columnspan=3, padx=10, pady=5)

tk.Button(form_frame, text="Salvar", command=enviar_dados, font=font_padrao, bg=cor_botao, fg="white", relief=tk.RAISED, bd=2).grid(row=5, column=1, columnspan=3, padx=10, pady=10)

# Adiciona o conteúdo da aba "Projetos Salvos"
frame_busca = tk.Frame(aba_projetos_salvos, padx=20, pady=20, bg=cor_aba)
frame_busca.pack(fill=tk.X)

tk.Label(frame_busca, text="Buscar Demandas:", font=font_padrao, bg=cor_aba, fg=cor_label).pack(side=tk.LEFT, padx=10, pady=5)
busca_entry = tk.Entry(frame_busca, width=60, font=font_padrao)
busca_entry.pack(side=tk.LEFT, padx=10, pady=5)
tk.Button(frame_busca, text="Buscar", command=buscar_demandas, font=font_padrao, bg="#FFC107", fg="black", relief=tk.RAISED, bd=2).pack(side=tk.LEFT, padx=10, pady=5)

resultado_text = scrolledtext.ScrolledText(aba_projetos_salvos, height=12, width=80, wrap=tk.WORD, font=font_padrao)
resultado_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

frame_projetos_container = tk.Frame(aba_projetos_salvos)
frame_projetos_container.pack(fill=tk.BOTH, expand=True)

canvas = tk.Canvas(frame_projetos_container)
scroll_y = tk.Scrollbar(frame_projetos_container, orient="vertical", command=canvas.yview)
projetos_frame = tk.Frame(canvas, bg=cor_aba)

projetos_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

canvas.create_window((0, 0), window=projetos_frame, anchor="nw")
canvas.configure(yscrollcommand=scroll_y.set)

scroll_y.pack(side="right", fill="y")
canvas.pack(side="left", fill="both", expand=True)

tk.Button(aba_projetos_salvos, text="Atualizar Projetos", command=atualizar_projetos_salvos, font=font_padrao, bg="#2196F3", fg="white", relief=tk.RAISED, bd=2).pack(padx=10, pady=10)

# Inicializa a lista de projetos salvos
atualizar_projetos_salvos()

# Inicia o loop principal
root.mainloop()
